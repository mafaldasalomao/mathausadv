from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx2pdf import convert
import tempfile
import math
import requests
import base64
import os
import shutil
from jinja2 import Template
import time


TEMP_FOLDER = 'temp_files'


tokenAPI = "live_8245d1d64e76f878f8aea00773d130afd7198ee9653700e1782c680dcb7aa3ea"
cryptKey = "live_crypt_pN5EBWOZnvybDnNF6mwJf4dJEXiaEFMn"


def preprocess_data(contratantes, responsaveis, service, fees):
    clients = contratantes
    if len(clients)>0:
        for index, item in enumerate(clients):
            item['title'] =  f'Contratante {index + 1}'
            item['representative'] = False
    else:
        for index, item in enumerate(clients):
            item['title'] =  f'Contratante'
            item['representative'] = False
    for index, item in enumerate(responsaveis):
        item['title'] =  f'Representante'
        item['representative'] = True
        clients.append(item)

    # Contexto a ser usado no template
    context = {
        "clients": clients,
        "service":service,
        "fees": fees
        
    }

    return context


def generate_contract(template_path, context):
    os.makedirs(TEMP_FOLDER, exist_ok=True)
    doc = Document(template_path)
    non_representatives = [person for person in context['clients'] if not person['representative']]
    for paragraph in doc.paragraphs:
        if "for person" in paragraph.text:
            template = Template(paragraph.text)
            rendered_text = template.render(context)
            paragraph.clear()
            
            for index, person in enumerate(context['clients']):
                title_run = paragraph.add_run(f"{person['title']}\n")
                title_run.font.name = "Urbanist"
                title_run.font.size = Pt(12)
                cpf_run = paragraph.add_run(f"CPF – {person['cpf']}\n")
                cpf_run.font.name = "Urbanist"
                cpf_run.font.size = Pt(8)

                name_run = paragraph.add_run(person['name'] + "\n")
                # name_run.bold = True  # Aplica o negrito
                name_run.font.name = "Urbanist SemiBold" 
                name_run.font.size = Pt(10)  

                address_run = paragraph.add_run(person['address'] + "\n\n")
                address_run.font.name = "Urbanist"
                address_run.font.size = Pt(10)  
                paragraph.paragraph_format.space_after = Pt(0)  
                paragraph.paragraph_format.space_before = Pt(0)
                break_line = math.ceil((len(context['clients'])/2))
                if (index + 1) % break_line == 0 and index < len(context['clients']) - 1:  
                    paragraph.add_run().add_break(WD_BREAK.COLUMN)

            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Alinhamento à esquerda

        if "for contact " in paragraph.text:
            template = Template(paragraph.text)
            rendered_text = template.render(context)
            paragraph.clear()
            
            for index, person in enumerate(non_representatives):
                title_run = paragraph.add_run(f"{person['title']} - Email – {person['email']}\n{person['title']} - Telefone – {person['phone']}\n")
                title_run.font.name = "Urbanist"
                title_run.font.size = Pt(8)
            # Formatação adicional (opcional)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Alinhamento à esquerda
        
        if "for signature " in paragraph.text:
            template = Template(paragraph.text)
            rendered_text = template.render(context)
            paragraph.clear()
            
            for index, person in enumerate(non_representatives):
                title_run = paragraph.add_run(f"__________________________________\n{person['name']}\n{person['title']}\n\n")
                title_run.font.name = "Urbanist"
                title_run.font.size = Pt(10)
                break_line = math.ceil((len(non_representatives)/2))            
                if (index + 1) % break_line == 0 and index < len(non_representatives) - 1:  
                    paragraph.add_run().add_break(WD_BREAK.COLUMN)
            
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
        if "service" in paragraph.text:
            paragraph.clear()
            service_run = paragraph.add_run(context['service'])
            service_run.font.name = "Urbanist"
            service_run.font.size = Pt(10)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

        if "fees" in paragraph.text:
            paragraph.clear()
            fees_run = paragraph.add_run(context['fees'])
            fees_run.font.name = "Urbanist"
            fees_run.font.size = Pt(10)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)

    # Salve o novo documento
    timestamp_unix = int(time.time())
    output_docx =f'temp_files/contrato.docx'
    doc.save(output_docx)
    convert(output_docx, output_docx.replace('.docx', '.pdf'))
    return output_docx.replace('.docx', '.pdf')












def send_document(file_path):
    if not os.path.exists(file_path):
        print(f"Erro: O arquivo {file_path} não foi encontrado.")
        return

    # URL do endpoint de upload
    url = "https://secure.d4sign.com.br/api/v1/documents/db39e124-2540-4f02-921b-8bbb756ba9bc/uploadbinary"
    # tokenAPI = "live_8245d1d64e76f878f8aea00773d130afd7198ee9653700e1782c680dcb7aa3ea"
    # cryptKey = "live_crypt_pN5EBWOZnvybDnNF6mwJf4dJEXiaEFMn"
    uuid_folder = "a879d30e-d8c7-4a5d-b77a-da0ca1d3bd33"

    # Lendo o arquivo e convertendo para Base64
    with open(file_path, 'rb') as file:
        encoded_file = base64.b64encode(file.read()).decode('utf-8')
    extension = os.path.splitext(file_path)[1].lower()
    mime_type = "application/pdf"
    # Montando o payload com o arquivo Base64 e uuid_folder
     # Montando o payload com o arquivo Base64 e uuid_folder
    payload = {
        "base64_binary_file": encoded_file,
        "mime_type": mime_type,  # Ou o tipo MIME do seu arquivo
        "name": file_path.split("/")[-1],  # Nome do arquivo, extraído do caminho
        "uuid_folder": uuid_folder
    }


    headers = {
        "accept": "application/json",
        "content-type": "application/json"
    }
        # Parâmetros da URL
    params = {
        "tokenAPI": tokenAPI,
        "cryptKey": cryptKey,
    }

    response = requests.post(url, json=payload, headers=headers, params=params)
    document_id = None
    if response.status_code == 200:
        print("Arquivo enviado com sucesso!")
        response_data = response.json()
        document_id = response_data.get("uuid") 
        print("Resposta:", response.json())
    else:
        print("Falha ao enviar o arquivo:", response.status_code)
        print("Resposta:", response.text)
    shutil.rmtree(TEMP_FOLDER)
    return document_id

# document_id = send_document("output1.pdf") #7d09a61b-3fff-4e22-9dc6-1e1559f8c91f
# print(document_id)


def cadastrar_signatarios(clients, document_id):
    # URL do endpoint de cadastro de signatários
    url = f"https://secure.d4sign.com.br/api/v1/documents/{document_id}/createlist"
    # tokenAPI = "live_8245d1d64e76f878f8aea00773d130afd7198ee9653700e1782c680dcb7aa3ea"
    # cryptKey = "live_crypt_pN5EBWOZnvybDnNF6mwJf4dJEXiaEFMn"
    # Montando o array de signatários
    signers = []

    for client in clients:
        if client["representative"]:  # Se for representante
            signers = [{
                "email": client["email"],  # Usando o email do representante
                "act": "1",  # Valor fixo ou pode ser modificado conforme a lógica
                "foreign": "0",  # Valor fixo, ajustar se necessário
                "certificadoicpbr": "0",  # Valor fixo, ajustar se necessário
                "assinatura_presencial": "0",  # Valor fixo, ajustar se necessário
                "docauth": "1",  # Valor fixo, ajustar se necessário
                "docauthandselfie": "1"  # Valor fixo, ajustar se necessário
            }]
            break
        else:
            signer = {
                "email": client["email"], 
                "act": "1", 
                "foreign": "0",
                "certificadoicpbr": "0",
                "assinatura_presencial": "0",
                "docauth": "1", 
                "docauthandselfie": "1"
            }
            signers.append(signer)

    # Montando o payload com os signatários
    payload = {
        "signers": signers
    }

    # Cabeçalhos da requisição
    headers = {
        "accept": "application/json",
        "content-type": "application/json"
    }

    # Parâmetros da URL
    params = {
        "tokenAPI": tokenAPI,
        "cryptKey": cryptKey,
    }

    # Enviando a requisição POST com o payload JSON
    response = requests.post(url, json=payload, headers=headers, params=params)

    # Verificando a resposta
    if response.status_code == 200:
        print("Signatários cadastrados com sucesso!")
        print("Resposta:", response.json())
    else:
        print("Falha ao cadastrar signatários:", response.status_code)
        print("Resposta:", response.text)



# Chamando a função com a lista de signatários
# cadastrar_signatarios(contratantes, document_id="7d09a61b-3fff-4e22-9dc6-1e1559f8c91f")



def send_document_to_signer(document_id, skip_email="false", workflow=None, message=None):
    """
    Envia um documento para assinatura através da API do D4Sign.

    :param document_id: UUID do documento que será enviado para assinatura.
    :param token_api: Token de autenticação da API.
    :param crypt_key: Chave criptográfica para autenticação.
    :param skip_email: Define se deve ou não enviar e-mail para o signatário (opcional, padrão é "false").
    :param workflow: Nome do fluxo de trabalho para rastreamento (opcional).
    :param message: Mensagem que será enviada junto com o documento (opcional).
    :return: Resposta da API.
    """
    # URL do endpoint
    url = f"https://secure.d4sign.com.br/api/v1/documents/{document_id}/sendtosigner?tokenAPI={tokenAPI}&cryptKey={cryptKey}"

    # Payload com os dados a serem enviados na requisição
    payload = {
        "skip_email": skip_email,
        "workflow": workflow,
        "message": message
    }

    # Cabeçalhos da requisição
    headers = {
        "accept": "application/json",  # Espera resposta em JSON
        "content-type": "application/json"  # Envia o corpo da requisição em formato JSON
    }

    # Enviando a requisição POST
    response = requests.post(url, json=payload, headers=headers)

    # Retornando a resposta da API
    return response.json()  # Retorna a resposta em formato JSON


# response = send_document_to_signer(
#         document_id="7d09a61b-3fff-4e22-9dc6-1e1559f8c91f",
#         token_api=tokenAPI,
#         crypt_key=cryptKey,
#         skip_email="false",
#         workflow="processo_assinatura",
#         message="Por favor, assine este documento."
# )

# print(response)  